import pandas as pd
import numpy as np
import os
from docx import Document
from docx.shared import Mm


def fill_empty_rows(row):
    row_names = list()
    hold: str = ""
    for (idx, cat) in row.items():
        current = str(idx)
        if current.find("Unnamed") == -1:
            hold = current
        else:
            current = hold
        row_names.append(current)
        # print(f'{current} : {cat}')

    return row_names


def write_to_text(categories, headers, contents):
    with open("first surec.txt", "w") as f:
        for (cat, header, content) in zip(categories, headers, contents):
            f.write(cat)
            if cat != '':
                f.write("\n---\n")
            f.write(header)
            f.write('\n')
            f.write(content)


def write_to_docx(categories, headers, contents, file_name):
    doc = Document()

    print('file working ' + file_name)

    doc.add_heading('Surec 1', 0)

    for (cat, header, content) in zip(categories, headers, contents):
        doc.add_heading(cat, level=1)
        doc.add_paragraph(header)
        doc.add_paragraph(content)

        doc.save(file_name+'.docx')


def foo(start, end, file_name, header_df, envanter):
    continuum = envanter.iloc[start:end, :]
    # print(continuum)
    categories = list()
    headers = list()
    contents = list()
    ctr = 0
    for col in continuum:
        # headers
        title = header_df.iloc[:, ctr][0]
        desc = header_df.iloc[:, ctr][1]
        if type(desc) == float:
            desc = ''
        h = title + '\n' + desc
        headers.append(h)
        ctr += 1
        col_series = continuum[col]
        # print(f'contents: {col_series.values}')
        s = ''
        col_name = str(col)
        if col_name.find('Unnamed') == -1:
            categories.append(col_name)
        else:
            categories.append('')
        for values in col_series.values:
            if type(values) == float:
                continue
            for value in values:
                if type(value) != float:
                    s += value
        s += '\n'
        contents.append(s)

    write_to_docx(categories, headers, contents, file_name)


if __name__ == '__main__':
    excel = pd.ExcelFile('sandport.xlsx')
    envanter = pd.read_excel(excel, 'Envanter')

    header_row = envanter.iloc[0]
    header_names = fill_empty_rows(header_row)

    # envanter = envanter.set_axis(header_names, axis='columns')
    surec_col = envanter.iloc[:, 3]

    surec_indexes = list()
    for (idx, val) in surec_col.items():
        if type(val) != float:
            surec_indexes.append((idx, val))

    # for (idx, val) in surec_indexes:
    #     print(f'{idx} -> {val}')

    header_df = envanter.iloc[:2, :]
    # pop the headers
    surec_indexes.pop(0)
    surec_indexes.pop(0)

    for i in range(len(surec_indexes)-1):
        start, vals = surec_indexes[i]
        end, _ = surec_indexes[i+1]
        # print(f'start: {start} -> val: {vals}')
        # print(f'end: {end} -> val: {vale}')
        val = vals.replace("/", "")
        val = val.replace("\n", "")
        print(val)
        file_name = f'{i+1}- ' + val
        foo(start, end, file_name, header_df, envanter)
        print('did that')

    # write_to_docx(categories, headers, contents)
